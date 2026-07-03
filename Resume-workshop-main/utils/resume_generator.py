# utils/resume_generator.py
# ─────────────────────────────────────────────
# Converts final resume text from Agent 4
# into a downloadable .docx file
# ─────────────────────────────────────────────

import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_resume_docx(resume_text: str) -> bytes:
    """
    Takes plain resume text and returns DOCX as bytes.
    Streamlit's download_button accepts bytes directly.
    """
    doc = Document()

    # ── Narrow page margins (standard for resumes) ──
    for section in doc.sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    lines = resume_text.strip().split("\n")

    for i, line in enumerate(lines):
        line = line.strip()

        # Skip empty lines — add blank spacer
        if not line:
            doc.add_paragraph("")
            continue

        # Detect section headers like ## SKILLS or SKILLS:
        is_header = (
            line.startswith("##")
            or line.isupper()
            or (len(line) < 40 and line.endswith(":"))
        )

        # First line = candidate name (largest text)
        if i == 0:
            para = doc.add_paragraph()
            run = para.add_run(line.replace("#", "").strip())
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif is_header:
            para = doc.add_paragraph()
            run = para.add_run(
                line.replace("##", "").replace(":", "").strip().upper()
            )
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x16, 0x21, 0x3E)
            para.paragraph_format.space_before = Pt(10)
            para.paragraph_format.space_after = Pt(2)

        # Bullet points
        elif line.startswith(("•", "-", "*")):
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(line.lstrip("•-* ").strip())
            run.font.size = Pt(10)

        # Regular text
        else:
            para = doc.add_paragraph()
            run = para.add_run(line.replace("**", "").strip())
            run.font.size = Pt(10)

    # ── Save to bytes for Streamlit download ──
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()