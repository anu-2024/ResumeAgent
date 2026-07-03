# utils/pdf_extractor.py
# ─────────────────────────────────────────────
# Extracts raw text from uploaded resume files
# Supports: PDF and DOCX formats
# ─────────────────────────────────────────────

import pdfplumber
import docx
import io


def extract_text_from_pdf(file) -> str:
    """
    Reads a PDF file and returns all text as a string.
    pdfplumber goes page by page and extracts text from each.
    """
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:  # some pages may be empty
                text += page_text + "\n"

    if not text.strip():
        raise ValueError(
            "PDF looks empty or is image-based. "
            "Please use a text-based PDF or DOCX."
        )
    return text.strip()


def extract_text_from_docx(file) -> str:
    """
    Reads a DOCX file and returns all text as a string.
    python-docx reads paragraph by paragraph.
    """
    # Streamlit uploads files as bytes
    # python-docx needs a file-like object — so we wrap in BytesIO
    if hasattr(file, "read"):
        doc = docx.Document(io.BytesIO(file.read()))
    else:
        doc = docx.Document(file)

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs).strip()


def extract_resume_text(uploaded_file) -> str:
    """
    Main function called by Streamlit.
    Auto-detects file type and calls the right extractor.
    """
    filename = uploaded_file.name.lower()

    if filename.endswith(".pdf"):
        return extract_text_from_pdf(uploaded_file)

    elif filename.endswith(".docx"):
        return extract_text_from_docx(uploaded_file)

    else:
        raise ValueError(
            f"Unsupported file: {filename}. "
            "Please upload PDF or DOCX only."
        )